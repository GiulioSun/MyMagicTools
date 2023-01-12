import streamlit as st
import pandas as pd
from docx import Document
import io
from pandas._libs.tslibs.timestamps import Timestamp

# App name
st.header(
    '''
    MagicTool
    '''
)

# Menu
option_menu = ['Home', 'Documentation', 'Stamp']
choice = st.sidebar.selectbox('Menu', option_menu)

if choice == 'Home':
    st.title(
        '''
        Home
        '''
    )


elif choice == 'Documentation':
    st.title(
        '''
        Documentation
        '''
    )
    buffer = io.BytesIO()

    st.markdown('''
        # Tips for correct use
        ## MagicStamp
        Generates word documents starting from a model with different data
        '''
                )

    st.image('image.png')

    st.markdown('''1. __Upload a file with the data table__, allowed formats [.xlsx, .xls, .csv] maximum size [200 
    MB]. - Column names must be enclosed in __<<field_name>>__. - The table must contain at least the 
    __<<filename>>__ column. 2. __Load a word template__, dynamic words must be enclosed in <<field_name>> 
    corresponding to the name of the column containing the data. 3. __Select the field__ that contains the name for 
    saving the file. 4.  Enter the __destination path__ example=> [/ Users / Desktop / example_fold /] 5. choose the 
    __file format__. 6. Press the __Gooo!__ button '''
                )

    example_data = pd.DataFrame(
        {
            '<<name>>': ['AAA', 'BBB', 'CCC'],
            '<<codezip>>': [1111, 2222, 3333],
            '<<email>>': ['aaa@test.com', 'bbb@test.com', 'ccc@test.com']
        })


    @st.cache
    def file_converter(df):
        return df.to_csv().encode('utf-8')


    csv = file_converter(example_data)

    st.download_button(
        label='Example table',
        data=csv,
        file_name='example_data.csv',
        mime="csv"
    )

elif choice == 'Stamp':
    st.title(
        '''
        Stamp
        '''
    )
    data: pd.DataFrame()
    doc: Document()

    # Title function
    st.title('''
        MagicStamp
        ''')
    # File importso
    file_up = st.file_uploader(
        'Upload the files!',
        type=['xlsx', 'xls', 'csv', 'docx'],
        accept_multiple_files=True
    )
    if file_up is not None:
        for i in file_up:
            # b_data = i.read()
            if i.name.split('.')[-1] == 'csv':
                data = pd.read_csv(i, sep=',')
                # Column name choice
                option = st.selectbox(
                    'Select column name document',
                    data.columns
                )

                check_show = st.checkbox('Show data')
                if check_show:
                    st.write(data.head(20))
                else:
                    pass

            elif i.name.split('.')[-1] == 'xlsx' or i.name.split('.')[-1] == 'xls':
                data = pd.read_excel(i)
                # Column name choice
                option = st.selectbox(
                    'Select column name document',
                    data.columns
                )
                check_show = st.checkbox('Show data')
                if check_show:
                    st.write(data.head(20))
                else:
                    pass

            elif i.name.split('.')[-1] == 'docx':
                doc = Document(i)

        # Path for save doc
        dest_path = st.text_input('Insert destination path')
        # Select option
        doc_format = st.selectbox('Select format', ['word', 'pdf'])
        # Button start
        button = st.button('Gooo!')

        if button == True and doc_format == 'word':
            with st.spinner('Working in progress'):
                labels = data.columns
                dc = data.to_dict('list')
                for j in range(len(data.iloc[:])):
                    for i in doc.paragraphs:
                        for w in labels:
                            if w in i.text and type(dc.get(w)[j]) == Timestamp:
                                i.text = i.text.replace(w, str(dc.get(w)[j].strftime('%d/%m/%Y')))
                            elif w in i.text:
                                i.text = i.text.replace(w, str(dc.get(w)[j]))

                    doc.save(f'{dest_path}{str(dc.get(option)[j])}.docx')
            st.success('Complete!')
